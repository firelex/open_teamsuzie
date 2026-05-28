from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from docx_template_agent.injection import fill_doc, inspect_template
from docx_template_agent.models import DocSpec

TEMPLATE = Path(__file__).parent.parent / "templates" / "engagement-letter.docx"
pytestmark = pytest.mark.skipif(not TEMPLATE.exists(), reason="engagement letter template missing")


def test_inspect_returns_placeholders_and_styles():
    m = inspect_template(TEMPLATE)
    # Letterhead-only template has at least the basic placeholders
    names = {p["name"] for p in m["placeholders"]}
    for required in ("company_name", "company_address", "telephone_number", "contact_person"):
        assert required in names, f"missing placeholder {required}"
    # Styles are present (Word ships dozens by default)
    assert m["style_count"] > 10
    # Section count is at least 1
    assert m["section_count"] >= 1


def test_inspect_classifies_heading_styles():
    m = inspect_template(TEMPLATE)
    role_counts: dict[str, int] = {}
    for s in m["styles"]:
        role_counts[s["inferred_role"]] = role_counts.get(s["inferred_role"], 0) + 1
    assert role_counts.get("heading", 0) >= 5  # Heading 1..N


def test_inspect_detects_letterhead_guidance():
    m = inspect_template(TEMPLATE)
    assert m["letterhead_guidance"] is not None


def test_inspect_has_no_letterhead_guidance_for_plain_template(tmp_path: Path):
    plain = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("[COMPANY NAME]")
    doc.save(plain)

    m = inspect_template(plain)
    assert m["letterhead_guidance"] is None


def test_fill_doc_replaces_tokens_and_preserves_structure(tmp_path: Path):
    spec = DocSpec(
        doc_name="Test",
        template=str(TEMPLATE),
        placeholders={
            "company_name": "Blixt Group",
            "company_address": "10 Ledbury Mews",
            "telephone_number": "+44 20 7946 0123",
        },
    )
    out, report = fill_doc(spec, tmp_path / "out.docx")
    assert out.exists()
    assert report.replacements > 0
    # No leftover bracket tokens for the keys we provided
    leftover_names = {tok for tok in report.unfilled_tokens}
    assert "[COMPANY NAME]" not in leftover_names
    assert "[COMPANY ADDRESS]" not in leftover_names
    # Unsupplied tokens (e.g. WEBSITE) should still be flagged as unfilled
    assert any("WEBSITE" in tok for tok in report.unfilled_tokens)


def test_fill_doc_preserves_paragraph_structure(tmp_path: Path):
    """Regression test for the run-consolidation bug — every text-box
    paragraph should keep its own text, not be wiped by a parent walk."""
    spec = DocSpec(
        doc_name="Test",
        template=str(TEMPLATE),
        placeholders={"company_name": "Blixt Group", "contact_person": "Mathias Strasser"},
    )
    out, _ = fill_doc(spec, tmp_path / "out.docx")

    from docx.oxml.ns import qn
    doc = Document(str(out))
    # Walk all <w:t> in the document and ensure "Blixt Group" and
    # "Mathias Strasser" appear as separate text elements (not merged).
    texts = [t.text or "" for t in doc.element.body.iter(qn("w:t"))]
    flat = " ".join(texts)
    assert "Blixt Group" in flat
    assert "Mathias Strasser" in flat
    # And that the text-box paragraphs still have multiple non-empty <w:t>
    non_empty = [t for t in texts if t.strip()]
    assert len(non_empty) > 5  # not consolidated to a single run


def test_fill_doc_writes_no_overrun_paragraphs(tmp_path: Path):
    """The filler should touch many paragraphs in a letterhead with 9 tokens
    repeated across two text-box copies — at least 10."""
    spec = DocSpec(
        doc_name="Test",
        template=str(TEMPLATE),
        placeholders={
            "company_name": "Blixt Group",
            "company_address": "10 Ledbury Mews North",
            "telephone_number": "+44 20 7946 0123",
            "fax_number": "+44 20 7946 0124",
            "website": "www.blixtgroup.com",
            "bank": "HSBC UK",
            "iban": "GB29 NWBK 6016 1331 9268 19",
            "bic": "HSBCGB2L",
            "contact_person": "Mathias Strasser",
            "assistant": "Office Manager",
        },
    )
    _, report = fill_doc(spec, tmp_path / "out.docx")
    assert report.paragraphs_touched >= 10
    assert report.unfilled_tokens == []

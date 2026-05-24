"""Anthropic-style tool definitions and dispatch for the docx agent loop.

The drafting flow is outline-first:

    inspect_template → [inspect_precedent] → set_placeholders →
    [append_paragraph for preamble] →
    set_outline → write_section ×N → save_doc

Each outline entry is a section *plan* with an internal `section_topic`
identifier, an optional rendered `section_heading` (null for sections
with no heading — e.g. a sign-off block), and an optional
`section_description` that helps the agent plan its content.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ..injection import (
    UnknownStyleError,
    append_paragraph,
    inspect_template,
    known_paragraph_styles,
    write_first_or_append,
)
from ..injection.filler import fill_doc
from ..injection.numbering import _is_list_style, restart_paragraph_numbering
from ..models import DocSpec
from ..services.precedents import inspect as inspect_precedent_service
from ..services.templates import get_template_path

DEFAULT_SECTION_HEADING_STYLE = "Heading 1"


TOOL_DEFINITIONS: list[dict[str, Any]] = [
    {
        "name": "inspect_template",
        "description": (
            "Return the template's structure: placeholders (bracketed "
            "tokens), styles (paragraph/character/table), sections, and "
            "outline. Always call this first."
        ),
        "input_schema": {"type": "object", "properties": {}, "required": []},
    },
    {
        "name": "inspect_precedent",
        "description": (
            "Return a precedent's full text, heading outline, and "
            "paragraph list. Use when reviewing the precedents listed in "
            "the user request to mirror their structure and tone."
        ),
        "input_schema": {
            "type": "object",
            "properties": {"precedent_id": {"type": "string"}},
            "required": ["precedent_id"],
        },
    },
    {
        "name": "set_placeholders",
        "description": (
            "Fill the template's bracket-token placeholders with values. "
            "Pass a dict keyed by the normalized placeholder names from "
            "the manifest (e.g. 'company_name'). Unknown names are ignored."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "values": {
                    "type": "object",
                    "additionalProperties": {"type": "string"},
                },
            },
            "required": ["values"],
        },
    },
    {
        "name": "set_outline",
        "description": (
            "Declare the ordered list of sections the document body will "
            "contain. Each section has a `section_topic` (internal id), "
            "an optional `section_description` (planning hint, not "
            "rendered), and an optional `section_heading` (the rendered "
            "Heading 1 text — pass null for sections that should appear "
            "WITHOUT a heading, e.g. an unlabelled sign-off block). "
            "save_doc rejects until every declared section has body "
            "content. Strings are also accepted (treated as topic = heading)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "sections": {
                    "type": "array",
                    "minItems": 1,
                    "items": {
                        "oneOf": [
                            {"type": "string"},
                            {
                                "type": "object",
                                "properties": {
                                    "section_topic": {"type": "string"},
                                    "section_description": {"type": "string"},
                                    "section_heading": {"type": ["string", "null"]},
                                },
                                "required": ["section_topic"],
                            },
                        ],
                    },
                },
                "heading_style": {
                    "type": "string",
                    "description": "Style used for section headings; default 'Heading 1'.",
                },
            },
            "required": ["sections"],
        },
    },
    {
        "name": "write_section",
        "description": (
            "Write one section identified by `section_topic`. If the outline "
            "entry has a non-null `section_heading`, it's emitted first. "
            "Then each entry in `paragraphs` is appended (each with `style` "
            "+ either `text` or `runs` for inline formatting). "
            "Strings are also accepted in `paragraphs` (become plain text)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "section_topic": {"type": "string"},
                "paragraphs": {
                    "type": "array",
                    "minItems": 1,
                    "items": {
                        "oneOf": [
                            {"type": "string"},
                            {
                                "type": "object",
                                "properties": {
                                    "style": {"type": "string"},
                                    "text": {"type": "string"},
                                    "runs": {
                                        "type": "array",
                                        "items": {
                                            "type": "object",
                                            "properties": {
                                                "text": {"type": "string"},
                                                "bold": {"type": "boolean"},
                                                "italic": {"type": "boolean"},
                                                "underline": {"type": "boolean"},
                                            },
                                            "required": ["text"],
                                        },
                                    },
                                },
                            },
                        ],
                    },
                },
            },
            "required": ["section_topic", "paragraphs"],
        },
    },
    {
        "name": "read_section",
        "description": (
            "Return the heading (if any) + body paragraphs already written "
            "for the section identified by `section_topic`. Use to maintain "
            "voice/coherence before writing or revising a neighbouring section."
        ),
        "input_schema": {
            "type": "object",
            "properties": {"section_topic": {"type": "string"}},
            "required": ["section_topic"],
        },
    },
    {
        "name": "revise_section",
        "description": (
            "Replace a section's body paragraphs in place. The heading (if "
            "any) stays; the body is removed and replaced with `paragraphs`."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "section_topic": {"type": "string"},
                "paragraphs": {
                    "type": "array",
                    "minItems": 1,
                    "items": {
                        "oneOf": [
                            {"type": "string"},
                            {"type": "object"},
                        ],
                    },
                },
            },
            "required": ["section_topic", "paragraphs"],
        },
    },
    {
        "name": "append_paragraph",
        "description": (
            "One-off paragraph append, OUTSIDE the section flow. Use for "
            "letter/memo preamble (confidentiality notices, addressee "
            "block, date, subject line, salutation, opening paragraph) "
            "BEFORE set_outline. Pass `text` for plain or `runs` for "
            "inline formatting; `style` must be a manifest style."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "text": {"type": "string"},
                "style": {"type": "string"},
                "runs": {"type": "array"},
            },
        },
    },
    {
        "name": "save_doc",
        "description": (
            "Validate completeness and write the .docx. Rejects the save "
            "if any section declared in set_outline lacks body content."
        ),
        "input_schema": {
            "type": "object",
            "properties": {"doc_name": {"type": "string"}},
            "required": ["doc_name"],
        },
    },
]


# ---- SectionPlan + AgentState ------------------------------------------------


@dataclass
class SectionPlan:
    """One outline entry: identifier + optional rendered heading + optional
    planning hint. After `write_section`, we also track lxml element refs
    for the heading and body paragraphs so read/revise can find them
    reliably even when the section has no heading text."""
    section_topic: str
    section_heading: str | None = None          # None = no heading rendered
    section_description: str | None = None
    written: bool = False
    heading_element: Any = None                 # <w:p> element of the heading, if any
    body_elements: list[Any] = field(default_factory=list)  # <w:p> elements of body paragraphs


@dataclass
class AgentState:
    template_path: Path
    output_path: Path
    document: Any = None
    placeholders: dict[str, str] = field(default_factory=dict)

    # Outline-first state
    outline: list[SectionPlan] = field(default_factory=list)
    heading_style: str = DEFAULT_SECTION_HEADING_STYLE

    # Bookkeeping
    appended_paragraphs: int = 0
    first_paragraph_consumed: bool = False
    last_paragraph_style: str | None = None
    list_restarts_applied: int = 0
    saved: bool = False
    save_report: dict[str, Any] | None = None

    def doc(self) -> Any:
        if self.document is None:
            self.document = Document(str(self.template_path))
        return self.document

    def section(self, topic: str) -> SectionPlan | None:
        for s in self.outline:
            if s.section_topic == topic:
                return s
        return None


# ---- Helpers -----------------------------------------------------------------


def _coerce_section_spec(raw: Any) -> SectionPlan:
    """Accept a string or a dict and return a SectionPlan."""
    if isinstance(raw, str):
        s = raw.strip()
        return SectionPlan(section_topic=s, section_heading=s)
    if isinstance(raw, dict):
        topic = (raw.get("section_topic") or raw.get("topic") or "").strip()
        if not topic:
            raise ValueError("each outline entry needs a non-empty section_topic")
        heading = raw.get("section_heading", raw.get("heading"))
        # Treat empty string heading as null (no heading)
        if isinstance(heading, str) and not heading.strip():
            heading = None
        description = raw.get("section_description") or raw.get("description")
        return SectionPlan(
            section_topic=topic,
            section_heading=heading,
            section_description=description,
        )
    raise ValueError(f"outline entry must be string or dict, got {type(raw).__name__}")


def _coerce_paragraph(para: Any) -> dict[str, Any]:
    """Accept either {style,text/runs} or a bare string."""
    if isinstance(para, str):
        return {"text": para, "style": None}
    if isinstance(para, dict):
        return para
    raise ValueError(f"paragraph must be a dict or string, got {type(para).__name__}")


def _maybe_restart_numbering(state: AgentState, doc, style: str | None) -> bool:
    """Apply the non-list→list numbering-restart rule."""
    curr_is_list = _is_list_style(style)
    prev_is_list = _is_list_style(state.last_paragraph_style)
    if curr_is_list and not prev_is_list:
        last_para = doc.paragraphs[-1]
        if restart_paragraph_numbering(doc, last_para, style):
            state.list_restarts_applied += 1
            state.last_paragraph_style = style
            return True
    state.last_paragraph_style = style
    return False


def _emit_paragraph(state: AgentState, doc, *, text=None, style=None, runs=None) -> Any:
    """Append (or seed-replace) a paragraph. Returns the `<w:p>` element of
    the resulting paragraph for downstream tracking."""
    if state.first_paragraph_consumed:
        append_paragraph(doc, text=text, style=style, runs=runs)
    else:
        write_first_or_append(doc, text=text, style=style, runs=runs)
        state.first_paragraph_consumed = True
    state.appended_paragraphs += 1
    _maybe_restart_numbering(state, doc, style)
    return doc.paragraphs[-1]._p


def _paragraph_to_spec(p) -> dict[str, Any]:
    """Serialise a python-docx Paragraph as a {style, text/runs} dict."""
    style = p.style.name if p.style else None
    runs = []
    for r in p.runs:
        rs: dict[str, Any] = {"text": r.text or ""}
        if r.bold:
            rs["bold"] = True
        if r.italic:
            rs["italic"] = True
        if r.underline:
            rs["underline"] = True
        runs.append(rs)
    if len(runs) == 1 and len(runs[0]) == 1:
        return {"style": style, "text": runs[0]["text"]}
    return {"style": style, "runs": runs}


def _paragraph_for_element(doc, element) -> Paragraph | None:
    for p in doc.paragraphs:
        if p._p is element:
            return p
    return None


# ---- Dispatch ----------------------------------------------------------------


def dispatch(state: AgentState, tool_name: str, tool_input: dict[str, Any]) -> Any:
    if tool_name == "inspect_template":
        return inspect_template(state.template_path)

    if tool_name == "inspect_precedent":
        try:
            return inspect_precedent_service(tool_input["precedent_id"])
        except FileNotFoundError as e:
            return {"error": str(e)}

    if tool_name == "set_placeholders":
        values = tool_input.get("values", {})
        state.placeholders.update(values)
        return {"ok": True, "total_placeholders": len(state.placeholders)}

    if tool_name == "set_outline":
        try:
            sections = [_coerce_section_spec(s) for s in tool_input.get("sections", [])]
        except ValueError as e:
            return {"error": str(e)}
        if not sections:
            return {"error": "outline must contain at least one section"}
        # Reject duplicate topics
        topics = [s.section_topic for s in sections]
        if len(set(topics)) != len(topics):
            return {"error": "section_topic values must be unique"}
        state.outline = sections
        state.heading_style = tool_input.get("heading_style") or DEFAULT_SECTION_HEADING_STYLE
        return {
            "ok": True,
            "heading_style": state.heading_style,
            "outline": [
                {
                    "section_topic": s.section_topic,
                    "section_heading": s.section_heading,
                    "section_description": s.section_description,
                }
                for s in state.outline
            ],
        }

    if tool_name == "write_section":
        topic = tool_input.get("section_topic") or tool_input.get("name") or ""
        topic = topic.strip()
        section = state.section(topic)
        if section is None:
            return {
                "error": f"unknown section_topic {topic!r}; declare it via set_outline first",
                "known_topics": [s.section_topic for s in state.outline],
            }
        if section.written:
            return {
                "error": f"section {topic!r} already written; use revise_section to replace its body",
            }
        paragraphs = tool_input.get("paragraphs", [])
        if not paragraphs:
            return {"error": "write_section requires at least one paragraph"}

        doc = state.doc()

        # Emit heading if requested
        if section.section_heading:
            try:
                section.heading_element = _emit_paragraph(
                    state, doc, text=section.section_heading, style=state.heading_style,
                )
            except UnknownStyleError as e:
                return {
                    "error": str(e),
                    "available_styles": sorted(known_paragraph_styles(doc)),
                }
            except ValueError as e:
                return {"error": str(e)}

        # Emit body paragraphs
        for i, raw in enumerate(paragraphs):
            try:
                para = _coerce_paragraph(raw)
            except ValueError as e:
                return {"error": str(e), "section_topic": topic, "paragraph_index": i}
            try:
                el = _emit_paragraph(
                    state, doc,
                    text=para.get("text"),
                    style=para.get("style"),
                    runs=para.get("runs"),
                )
            except UnknownStyleError as e:
                return {
                    "error": str(e), "section_topic": topic, "paragraph_index": i,
                    "available_styles": sorted(known_paragraph_styles(doc)),
                }
            except ValueError as e:
                return {"error": str(e), "section_topic": topic, "paragraph_index": i}
            section.body_elements.append(el)

        section.written = True
        return {
            "ok": True,
            "section_topic": topic,
            "heading_emitted": section.section_heading is not None,
            "body_paragraph_count": len(section.body_elements),
            "sections_written": [s.section_topic for s in state.outline if s.written],
        }

    if tool_name == "read_section":
        topic = (tool_input.get("section_topic") or tool_input.get("name") or "").strip()
        section = state.section(topic)
        if section is None:
            return {
                "error": f"unknown section_topic {topic!r}",
                "known_topics": [s.section_topic for s in state.outline],
            }
        if not section.written:
            return {"error": f"section {topic!r} has not been written yet"}
        doc = state.doc()
        heading_spec = None
        if section.heading_element is not None:
            hp = _paragraph_for_element(doc, section.heading_element)
            if hp is not None:
                heading_spec = _paragraph_to_spec(hp)
        body = []
        for el in section.body_elements:
            bp = _paragraph_for_element(doc, el)
            if bp is not None:
                body.append(_paragraph_to_spec(bp))
        return {
            "section_topic": topic,
            "section_heading": section.section_heading,
            "heading": heading_spec,
            "body": body,
            "body_paragraph_count": len(body),
        }

    if tool_name == "revise_section":
        topic = (tool_input.get("section_topic") or tool_input.get("name") or "").strip()
        section = state.section(topic)
        if section is None:
            return {
                "error": f"unknown section_topic {topic!r}",
                "known_topics": [s.section_topic for s in state.outline],
            }
        if not section.written:
            return {"error": f"section {topic!r} has not been written yet; use write_section"}
        paragraphs = tool_input.get("paragraphs", [])
        if not paragraphs:
            return {"error": "revise_section requires at least one paragraph"}
        doc = state.doc()

        # Choose an anchor: the section's heading element (if any) is the
        # insertion point; otherwise the first body element marks the
        # start. New paragraphs go immediately after the anchor.
        if section.heading_element is not None:
            anchor = section.heading_element
        elif section.body_elements:
            anchor = section.body_elements[0].getprevious()
        else:
            anchor = None

        # Remove existing body
        for el in section.body_elements:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        section.body_elements = []

        # Reset numbering state to the heading style so list-restart logic
        # treats this as a fresh non-list → list transition.
        state.last_paragraph_style = state.heading_style if section.heading_element is not None else None

        for i, raw in enumerate(paragraphs):
            try:
                para = _coerce_paragraph(raw)
            except ValueError as e:
                return {"error": str(e), "section_topic": topic, "paragraph_index": i}
            try:
                append_paragraph(
                    doc,
                    text=para.get("text"),
                    style=para.get("style"),
                    runs=para.get("runs"),
                )
            except UnknownStyleError as e:
                return {
                    "error": str(e), "section_topic": topic, "paragraph_index": i,
                    "available_styles": sorted(known_paragraph_styles(doc)),
                }
            new_p = doc.paragraphs[-1]._p
            if anchor is not None:
                anchor.addnext(new_p)
                anchor = new_p
            section.body_elements.append(new_p)
            state.appended_paragraphs += 1
            _maybe_restart_numbering(state, doc, para.get("style"))

        return {
            "ok": True,
            "section_topic": topic,
            "body_paragraph_count": len(section.body_elements),
        }

    if tool_name == "append_paragraph":
        text = tool_input.get("text")
        style = tool_input.get("style")
        runs = tool_input.get("runs")
        doc = state.doc()
        try:
            _emit_paragraph(state, doc, text=text, style=style, runs=runs)
        except UnknownStyleError as e:
            return {
                "error": str(e),
                "available_styles": sorted(known_paragraph_styles(doc)),
            }
        except ValueError as e:
            return {"error": str(e)}
        return {"ok": True, "appended": state.appended_paragraphs}

    if tool_name == "save_doc":
        if state.outline:
            missing = [s.section_topic for s in state.outline if not s.written]
            if missing:
                return {
                    "error": "save refused: declared sections missing body content",
                    "missing_topics": missing,
                    "guidance": (
                        "Call write_section for each missing topic before "
                        "saving."
                    ),
                }
            empty = [
                s.section_topic for s in state.outline
                if not s.body_elements
            ]
            if empty:
                return {
                    "error": "save refused: declared sections have no body",
                    "empty_topics": empty,
                }

        state.doc().save(str(state.output_path))
        spec = DocSpec(
            doc_name=tool_input["doc_name"],
            template=str(state.output_path),
            placeholders=state.placeholders,
        )
        path, report = fill_doc(spec, state.output_path)
        state.saved = True
        state.save_report = {
            "path": str(path),
            "replacements": report.replacements,
            "paragraphs_touched": report.paragraphs_touched,
            "paragraphs_pruned": report.paragraphs_pruned,
            "paragraphs_appended": state.appended_paragraphs,
            "list_restarts_applied": state.list_restarts_applied,
            "outline": [
                {
                    "section_topic": s.section_topic,
                    "section_heading": s.section_heading,
                    "written": s.written,
                }
                for s in state.outline
            ],
            "unfilled_tokens": report.unfilled_tokens,
        }
        return state.save_report

    raise ValueError(f"Unknown tool: {tool_name}")


def state_from_template_id(template_id: str, output_path: Path) -> AgentState:
    return AgentState(
        template_path=get_template_path(template_id),
        output_path=output_path,
    )

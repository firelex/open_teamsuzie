"""Append styled paragraphs to the body of a docx Document.

Used by the LLM agent loop when drafting content into a letterhead-only
template — the letterhead lives in headers/footers or text-box drawings,
the body accepts ordinary styled paragraphs."""

from __future__ import annotations

from docx.document import Document
from docx.oxml.ns import qn
from typing import TypedDict


class UnknownStyleError(ValueError):
    """The requested style isn't defined in the template."""


class RunSpec(TypedDict, total=False):
    text: str
    bold: bool
    italic: bool
    underline: bool


def known_paragraph_styles(doc: Document) -> set[str]:
    names: set[str] = set()
    for s in doc.styles:
        try:
            stype = str(s.type)
        except Exception:
            stype = ""
        if "PARAGRAPH" in stype.upper():
            names.add(s.name)
    return names


def _first_empty_body_paragraph(doc: Document):
    """Return the first body paragraph that has no direct text content.

    Used to detect the template's empty seed paragraph at the top of the
    body (often holding the letterhead drawing as an anchor). On the first
    `append_paragraph` call we write into this paragraph instead of adding
    a new one, so the output doesn't start with a phantom blank line."""
    for p in doc.paragraphs:
        # Direct text only — skip nested-paragraph text from drawings/text-boxes
        direct_text = []
        for t in p._p.iter(qn("w:t")):
            parent_p = t.getparent()
            while parent_p is not None and parent_p.tag != qn("w:p"):
                parent_p = parent_p.getparent()
            if parent_p is p._p:
                direct_text.append(t.text or "")
        if not "".join(direct_text).strip():
            return p
        return None  # the first body paragraph has text — don't reuse
    return None


def _apply_runs_to_paragraph(para, text: str | None, runs: list[RunSpec] | None) -> None:
    if runs:
        for r in runs:
            run = para.add_run(r.get("text", ""))
            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            if r.get("underline"):
                run.underline = True
    else:
        para.add_run(text or "")


def append_paragraph(
    doc: Document,
    text: str | None = None,
    style: str | None = None,
    runs: list[RunSpec] | None = None,
) -> None:
    """Add a new body paragraph.

    Pass either `text` for a single-run paragraph, or `runs` for inline
    formatting (e.g. bold lead phrase: `[{"text": "Strong experience:",
    "bold": True}, {"text": " Our team has 100+ years..."}]`).

    When `style` is given it must exist in the template's style table —
    refusing to invent style names is the contract that keeps the output
    on house-style."""
    if text is None and not runs:
        raise ValueError("append_paragraph requires either `text` or `runs`")
    if style is not None and style not in known_paragraph_styles(doc):
        raise UnknownStyleError(
            f"style {style!r} is not defined on this template; "
            "use a style name returned by inspect_template"
        )

    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _apply_runs_to_paragraph(para, text, runs)


def write_first_or_append(
    doc: Document,
    text: str | None = None,
    style: str | None = None,
    runs: list[RunSpec] | None = None,
) -> bool:
    """If the template's first body paragraph is text-empty, write content
    into it (preserving any drawing-carrier runs the template put there for
    letterhead positioning) and return True. Otherwise behave like
    `append_paragraph` and return False.

    The agent loop uses this for the very first paragraph it emits, so the
    output doesn't start with a phantom blank line."""
    if text is None and not runs:
        raise ValueError("write_first_or_append requires either `text` or `runs`")
    if style is not None and style not in known_paragraph_styles(doc):
        raise UnknownStyleError(
            f"style {style!r} is not defined on this template; "
            "use a style name returned by inspect_template"
        )

    seed = _first_empty_body_paragraph(doc)
    if seed is None:
        append_paragraph(doc, text=text, style=style, runs=runs)
        return False

    if style:
        seed.style = doc.styles[style]
    _apply_runs_to_paragraph(seed, text, runs)
    return True


def prune_empty_paragraphs(doc: Document) -> int:
    """Remove every body paragraph whose text is empty or whitespace-only.

    Lets the agent rely on the paragraph styles' built-in space-after
    rather than inserting blank spacer paragraphs manually. Skips
    paragraphs inside tables, headers, footers, and drawings — only
    direct body-level paragraphs are pruned."""
    removed = 0
    body = doc.element.body
    for p in list(body.findall(qn("w:p"))):
        # Concatenate direct text (avoid nested drawing text)
        texts = []
        for t in p.iter(qn("w:t")):
            # Skip text inside nested paragraphs (text boxes)
            parent_p = t.getparent()
            while parent_p is not None and parent_p.tag != qn("w:p"):
                parent_p = parent_p.getparent()
            if parent_p is p:
                texts.append(t.text or "")
        full = "".join(texts).strip()
        if not full:
            # Don't strip paragraphs that carry a drawing (letterhead in a
            # text-box) or nested paragraphs (text-box content) — only
            # genuinely empty body paragraphs are spacer artefacts.
            has_drawing = next(p.iter(qn("w:drawing")), None) is not None
            has_nested_p = any(
                child is not p for child in p.iter(qn("w:p"))
            )
            if has_drawing or has_nested_p:
                continue
            body.remove(p)
            removed += 1
    return removed

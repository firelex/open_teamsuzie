"""Template introspection: walk a .docx and emit a manifest of styles,
placeholders, sections.

The manifest is intentionally permissive — sections, outline, content
controls and "sample" body content are all optional fields. A template
that ships as letterhead-only (one empty body paragraph, populated
headers/footers) still produces a usable manifest."""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path
from typing import TypedDict

from docx import Document
from docx.oxml.ns import qn

from .tokens import find_tokens_in, normalize_token

WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
VML_NS = "urn:schemas-microsoft-com:vml"
LETTERHEAD_GUIDANCE = (
    "This template has a floating letterhead text box. "
    "Substantive body content should start below the float, using "
    "a Word clear-wrap break before the first body paragraph."
)


class StyleManifest(TypedDict):
    name: str
    type: str                   # paragraph | character | table | numbering
    font_name: str | None
    font_size_pt: float | None
    bold: bool | None
    italic: bool | None
    color: str | None
    based_on: str | None
    inferred_role: str          # "heading" | "body" | "caption" | "list" | "other"


class PlaceholderManifest(TypedDict):
    name: str                   # normalized
    token: str                  # raw bracketed form
    location: str               # body | header | footer
    section_index: int
    sample_paragraph: str
    paragraph_style: str | None


class SectionManifest(TypedDict):
    index: int
    page_width_in: float
    page_height_in: float
    margin_left_in: float
    margin_right_in: float
    margin_top_in: float
    margin_bottom_in: float
    has_header: bool
    has_footer: bool


class TemplateManifest(TypedDict):
    paragraph_count: int
    style_count: int
    section_count: int
    placeholder_count: int
    styles: list[StyleManifest]
    placeholders: list[PlaceholderManifest]
    sections: list[SectionManifest]
    outline: list[str]          # heading texts in document order; empty if none
    letterhead_guidance: str | None


def _infer_role(style_name: str, style_type: str) -> str:
    n = style_name.lower()
    is_para = "paragraph" in style_type
    if is_para:
        if "heading" in n or n == "title":
            return "heading"
        if "list" in n or "bullet" in n or n.startswith("number"):
            return "list"
        if "caption" in n:
            return "caption"
        if "quote" in n:
            return "quote"
        if n in {"normal", "body", "body text"}:
            return "body"
    return "other"


def _walk_paragraph_text(part_paragraphs) -> Iterable[tuple[str, str | None]]:
    """Yield (text, style_name) for each paragraph in a container."""
    for p in part_paragraphs:
        yield p.text, (p.style.name if p.style else None)


def _walk_xml_paragraphs(part_element) -> Iterable[str]:
    """Walk every `<w:p>` descendant under `part_element`, returning the
    concatenated text. Catches paragraphs inside tables, content controls,
    and other containers python-docx's `.paragraphs` collection misses."""
    for p in part_element.iter(qn("w:p")):
        texts = [t.text or "" for t in p.iter(qn("w:t"))]
        yield "".join(texts)


def _styles(doc: Document) -> list[StyleManifest]:
    out: list[StyleManifest] = []
    for style in doc.styles:
        try:
            name = style.name
            stype = str(style.type).split(".")[-1].lower() if style.type else "unknown"
        except Exception:
            continue
        font = getattr(style, "font", None)
        out.append(StyleManifest(
            name=name,
            type=stype,
            font_name=getattr(font, "name", None) if font is not None else None,
            font_size_pt=(font.size.pt if font is not None and font.size is not None else None),
            bold=getattr(font, "bold", None) if font is not None else None,
            italic=getattr(font, "italic", None) if font is not None else None,
            color=(
                font.color.rgb.__str__() if (
                    font is not None and font.color is not None
                    and font.color.rgb is not None
                ) else None
            ),
            based_on=(style.base_style.name if getattr(style, "base_style", None) else None),
            inferred_role=_infer_role(name, stype),
        ))
    return out


def _sections(doc: Document) -> list[SectionManifest]:
    out: list[SectionManifest] = []
    for i, s in enumerate(doc.sections):
        out.append(SectionManifest(
            index=i,
            page_width_in=(s.page_width.inches if s.page_width else 0),
            page_height_in=(s.page_height.inches if s.page_height else 0),
            margin_left_in=(s.left_margin.inches if s.left_margin else 0),
            margin_right_in=(s.right_margin.inches if s.right_margin else 0),
            margin_top_in=(s.top_margin.inches if s.top_margin else 0),
            margin_bottom_in=(s.bottom_margin.inches if s.bottom_margin else 0),
            has_header=any(p.text.strip() for p in s.header.paragraphs),
            has_footer=any(p.text.strip() for p in s.footer.paragraphs),
        ))
    return out


def _placeholders(doc: Document) -> list[PlaceholderManifest]:
    """Scan body + headers + footers (all sections) for bracket tokens."""
    found: dict[str, PlaceholderManifest] = {}

    def _record(token: str, location: str, section_index: int, sample: str, style: str | None) -> None:
        name = normalize_token(token)
        if name in found:
            return
        found[name] = PlaceholderManifest(
            name=name, token=token, location=location,
            section_index=section_index, sample_paragraph=sample[:200],
            paragraph_style=style,
        )

    # Body — walk via XML so tokens inside tables / content controls are caught
    for text in _walk_xml_paragraphs(doc.element.body):
        for tok in find_tokens_in(text):
            _record(tok, "body", 0, text, None)
    # Resolve paragraph styles for tokens found at the top level
    for p in doc.paragraphs:
        if not p.text:
            continue
        for tok in find_tokens_in(p.text):
            name = normalize_token(tok)
            if name in found and found[name]["paragraph_style"] is None:
                found[name]["paragraph_style"] = p.style.name if p.style else None

    for i, section in enumerate(doc.sections):
        for area, kind in (
            (section.header, "header"),
            (section.footer, "footer"),
            (section.first_page_header, "header"),
            (section.first_page_footer, "footer"),
        ):
            if area is None:
                continue
            for p in area.paragraphs:
                for tok in find_tokens_in(p.text):
                    _record(tok, kind, i, p.text, p.style.name if p.style else None)

    return list(found.values())


def _outline(doc: Document) -> list[str]:
    out: list[str] = []
    for p in doc.paragraphs:
        sname = (p.style.name if p.style else "") or ""
        if sname.lower().startswith("heading") or sname.lower() == "title":
            text = p.text.strip()
            if text:
                out.append(text)
    return out


def _letterhead_guidance(doc: Document) -> str | None:
    """Detect floating letterhead/text-box layouts that body content must clear."""

    wrap_tags = (qn("wp:wrapSquare"), qn("wp:wrapTight"))
    body_pr_tag = f"{{{WPS_NS}}}bodyPr"
    shape_tag = f"{{{VML_NS}}}shape"

    for anchor in doc.element.body.iter(qn("wp:anchor")):
        if any(anchor.find(f".//{tag}") is not None for tag in wrap_tags):
            return LETTERHEAD_GUIDANCE

        for body_pr in anchor.iter(body_pr_tag):
            if (body_pr.get("wrap") or "").lower() in {"square", "tight"}:
                return LETTERHEAD_GUIDANCE

    for shape in doc.element.body.iter(shape_tag):
        style = (shape.get("style") or "").lower()
        if "mso-wrap-style:square" in style or "mso-wrap-style:tight" in style:
            return LETTERHEAD_GUIDANCE

    return None


def inspect_template(template_path: str | Path) -> TemplateManifest:
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")

    doc = Document(str(path))
    styles = _styles(doc)
    sections = _sections(doc)
    placeholders = _placeholders(doc)
    outline = _outline(doc)
    letterhead_guidance = _letterhead_guidance(doc)

    # Count paragraphs across the whole document (not just body.paragraphs)
    total_paragraphs = sum(1 for _ in doc.element.body.iter(qn("w:p")))

    return TemplateManifest(
        paragraph_count=total_paragraphs,
        style_count=len(styles),
        section_count=len(sections),
        placeholder_count=len(placeholders),
        styles=styles,
        placeholders=placeholders,
        sections=sections,
        outline=outline,
        letterhead_guidance=letterhead_guidance,
    )
